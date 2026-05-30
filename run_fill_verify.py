# -*- coding: utf-8 -*-
"""Create template if needed, run fill_apap, verify output."""
from __future__ import annotations

import json
import subprocess
import sys
from pathlib import Path

OUT = Path(r"C:\Users\vlvlk\Desktop\апап.docx")
RESULT = Path(r"C:\Users\vlvlk\Desktop\fill_verify_result.json")
SCRIPT_DIR = Path(__file__).resolve().parent


def ensure_template() -> None:
    if OUT.exists():
        return
    try:
        from docx import Document
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
        from docx import Document

    doc = Document()
    doc.add_paragraph("")
    doc.save(str(OUT))


def verify_content() -> dict:
    from docx import Document

    doc = Document(str(OUT))
    text = "\n".join(p.text for p in doc.paragraphs)
    checks = {
        "ГЛАВА 1": "ГЛАВА 1" in text.upper(),
        "ГЛАВА 2": "ГЛАВА 2" in text.upper(),
        "ГЛАВА 3": "ГЛАВА 3" in text.upper(),
        "ЗАКЛЮЧЕНИЕ": "ЗАКЛЮЧЕНИЕ" in text.upper(),
        "оргтехник": "оргтехник" in text.lower(),
        "сайт": "сайт" in text.lower(),
    }
    return checks


def main() -> int:
    result: dict = {"success": False, "path": str(OUT)}
    try:
        ensure_template()
        r = subprocess.run(
            [sys.executable, str(SCRIPT_DIR / "fill_apap.py")],
            cwd=str(SCRIPT_DIR),
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
        )
        result["fill_stdout"] = r.stdout
        result["fill_stderr"] = r.stderr
        result["fill_returncode"] = r.returncode

        if not OUT.exists():
            result["error"] = "Output file not created"
            RESULT.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
            return 1

        size = OUT.stat().st_size
        result["size_bytes"] = size
        result["size_ok"] = size > 10000
        result["content_checks"] = verify_content()
        result["content_ok"] = all(result["content_checks"].values())
        result["success"] = (
            r.returncode == 0 and result["size_ok"] and result["content_ok"]
        )
        RESULT.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
        return 0 if result["success"] else 1
    except Exception as e:
        result["error"] = str(e)
        RESULT.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
